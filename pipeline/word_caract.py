"""
╔══════════════════════════════════════════════════════════════════════════════╗
║   SCRIPT_TOP_Universal_Word_Caracterizacion.py  —  v1.0                   ║
║   Genera informe Word de caracterización al ingreso (TOP1)                 ║
║   10 gráficos · Compatible con cualquier país TOP                          ║
╠══════════════════════════════════════════════════════════════════════════════╣
║  CÓMO USAR:                                                                 ║
║  1. Abre un chat nuevo con Claude                                           ║
║  2. Sube DOS archivos:                                                      ║
║       • Este script                                                         ║
║       • La base Wide (generada por SCRIPT_TOP_Universal_Wide)              ║
║  3. Escribe: "Ejecuta el script Word Caracterización TOP con esta base"    ║
╚══════════════════════════════════════════════════════════════════════════════╝
"""

import glob, os, unicodedata, io, warnings
import pandas as pd
import numpy as np
import matplotlib; matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
warnings.filterwarnings('ignore')

# ── Colores ───────────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x1F,0x38,0x64); BLUE  = RGBColor(0x2E,0x75,0xB6)
ACCENT= RGBColor(0x00,0xB0,0xF0); GRAY  = RGBColor(0x59,0x59,0x59)
WHITE = RGBColor(0xFF,0xFF,0xFF); LIGHT = RGBColor(0xBD,0xD7,0xEE)
C_NAVY='1F3864'; C_MID='2E75B6'; C_LIGHT='EEF4FB'; C_WHITE='FFFFFF'
MC_MID='#2E75B6'; MC_LIGHT='#BDD7EE'; MC_ACCENT='#00B0F0'; MC_GRAY='#BFBFBF'
PIE_COLS=['#2E75B6','#1F3864','#4472C4','#9DC3E6','#00B0F0','#538135','#BFBFBF','#C00000','#ED7D31']

# ── Detección de país ─────────────────────────────────────────────────────────
_PAISES = {
    'republica_dominicana':'República Dominicana','dominicana':'República Dominicana',
    'honduras':'Honduras','panama':'Panamá','panam':'Panamá',
    'el_salvador':'El Salvador','salvador':'El Salvador',
    'mexico':'México','mexic':'México','ecuador':'Ecuador',
    'peru':'Perú','argentina':'Argentina','colombia':'Colombia',
    'chile':'Chile','bolivia':'Bolivia','paraguay':'Paraguay',
    'uruguay':'Uruguay','venezuela':'Venezuela','guatemala':'Guatemala',
    'costa_rica':'Costa Rica','costarica':'Costa Rica','nicaragua':'Nicaragua',
}

def _norm(s):
    return unicodedata.normalize('NFD',str(s).lower()).encode('ascii','ignore').decode()

def _extraer_pais(fn):
    f = _norm(str(fn).replace('.','_'))
    for k,v in _PAISES.items():
        if k in f: return v
    return None

def _detectar_pais(wide_file):
    try:
        rs = pd.read_excel(wide_file, sheet_name='Resumen', header=None)
        for _, row in rs.iterrows():
            for v in row.tolist():
                p = _extraer_pais(str(v))
                if p: return p
    except: pass
    return _extraer_pais(os.path.basename(wide_file))

def auto_archivo_wide():
    candidatos = (
        glob.glob('/mnt/user-data/uploads/*Wide*.xlsx') +
        glob.glob('/mnt/user-data/uploads/*wide*.xlsx') +
        glob.glob('/mnt/user-data/uploads/TOP_Base*.xlsx') +
        glob.glob('/home/claude/TOP_Base_Wide.xlsx'))
    if not candidatos:
        raise FileNotFoundError('No se encontró la base Wide TOP.')
    print(f'  → Base Wide: {os.path.basename(candidatos[0])}')
    return candidatos[0]

# ══════════════════════════════════════════════════════════════════════════════
# CONFIGURACIÓN — inyectada por runner.py (no ejecutar a nivel módulo)
# ══════════════════════════════════════════════════════════════════════════════
INPUT_FILE    = None   # runner inyecta la ruta real
SHEET_NAME    = 'Base Wide'
OUTPUT_FILE   = None   # runner inyecta la ruta real
FILTRO_CENTRO = None   # runner inyecta el filtro si aplica
NOMBRE_SERVICIO = 'Servicio de Tratamiento'
PERIODO         = ''

# ══════════════════════════════════════════════════════════════════════════════
# HELPERS WORD
# ══════════════════════════════════════════════════════════════════════════════
def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'), hex_color); tcPr.append(shd)

def add_section_header(doc, num, title):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    c = tbl.rows[0].cells[0]; set_cell_bg(c, C_MID)
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    txt = f'  {num}. {title.upper()}' if num else f'  {title.upper()}'
    run = p.add_run(txt)
    run.font.name='Calibri'; run.font.size=Pt(11)
    run.font.bold=True; run.font.color.rgb=WHITE
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(0)
    _keep_with_next(sp)

def _keep_with_next(p):
    """Evita que el párrafo quede separado del siguiente elemento (gráfico/tabla)."""
    pPr = p._p.get_or_add_pPr()
    kwn = OxmlElement('w:keepNext')
    pPr.append(kwn)
    kl = OxmlElement('w:keepLines')
    pPr.append(kl)

def add_subsection(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(title); run.bold=True
    run.font.name='Calibri'; run.font.size=Pt(10.5)
    run.font.color.rgb = NAVY
    _keep_with_next(p)  # El título siempre acompaña al gráfico siguiente

def add_body(doc, text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name='Calibri'; run.font.size=Pt(10)
    run.font.italic=italic; run.font.color.rgb=RGBColor(0x33,0x33,0x33)
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name='Calibri'; run.font.size=Pt(8.5)
    run.font.italic=True; run.font.color.rgb=GRAY

def add_kpi_row(doc, kpis):
    tbl = doc.add_table(rows=2, cols=len(kpis))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j,(val,lbl) in enumerate(kpis):
        cv = tbl.rows[0].cells[j]; set_cell_bg(cv, C_LIGHT)
        p = cv.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(val))
        run.font.name='Calibri'; run.font.size=Pt(22)
        run.font.bold=True; run.font.color.rgb=BLUE
        cl = tbl.rows[1].cells[j]; set_cell_bg(cl, C_LIGHT)
        p2 = cl.paragraphs[0]; p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(lbl)
        run2.font.name='Calibri'; run2.font.size=Pt(8.5)
        run2.font.color.rgb=GRAY
    doc.add_paragraph()

def fig_to_img(fig, width_cm=12):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    buf.seek(0); plt.close(fig)
    return buf, Cm(width_cm)

def add_picture_kwnext(doc, buf, width):
    """Inserta imagen centrada con keepNext para que el texto descriptivo no se separe."""
    doc.add_picture(buf, width=width)
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _keep_with_next(p)

def _ax_style(ax, horiz=False):
    (ax.xaxis if horiz else ax.yaxis).grid(True, color='#E2E8F0', linewidth=0.6, zorder=0)
    ax.set_axisbelow(True)
    for sp in ['top','right']: ax.spines[sp].set_visible(False)
    ax.spines['left'].set_color('#D0D0D0'); ax.spines['bottom'].set_color('#D0D0D0')
    ax.set_facecolor('white')

# ══════════════════════════════════════════════════════════════════════════════
# DETECCIÓN DE COLUMNAS
# ══════════════════════════════════════════════════════════════════════════════
def detectar_columnas(cols):
    sust_cols = []
    for c in cols:
        if c.endswith('_TOP1') and 'Total (0-28)' in c:
            base = c.replace('_TOP1','')
            if base.startswith('1)'):
                partes = base.split('>>')
                if len(partes)>=3:
                    nombre = partes[-2].strip().split('(')[0].strip()
                    sust_cols.append((nombre,c))
    tr_sn = []
    for c in cols:
        if c.endswith('_TOP1') and c.replace('_TOP1','').startswith('3)') and '>>' in c:
            nombre = c.replace('_TOP1','').split('>>')[-1].strip()
            tr_sn.append((nombre,c))
    vif     = next((c for c in cols if c.endswith('_TOP1') and '4)' in c
                    and 'Violencia Intrafamiliar' in c and 'Total (0-28)' in c), None)
    sal_psi = next((c for c in cols if c.endswith('_TOP1') and c.replace('_TOP1','').startswith('6)')), None)
    sal_fis = next((c for c in cols if c.endswith('_TOP1') and c.replace('_TOP1','').startswith('8)')), None)
    cal_vid = next((c for c in cols if c.endswith('_TOP1') and c.replace('_TOP1','').startswith('10)')), None)
    viv1    = next((c for c in cols if c.endswith('_TOP1') and '9)' in c and 'estable' in c.lower()), None)
    viv2    = next((c for c in cols if c.endswith('_TOP1') and '9)' in c and 'básicas' in c.lower()), None)
    sust_pp = next((c for c in cols if c.endswith('_TOP1') and c.replace('_TOP1','').startswith('2)')
                    and 'sustancia principal' in c.lower()), None)
    sexo    = next((c for c in cols if c.endswith('_TOP1') and 'sexo' in c.lower()), None)
    fn_col  = next((c for c in cols if c.endswith('_TOP1') and 'nacimiento' in c.lower()), None)
    fecha   = next((c for c in cols if c.endswith('_TOP1') and 'fecha entrevista' in c.lower()), None)
    return dict(sust_cols=sust_cols,tr_sn=tr_sn,vif=vif,
                sal_psi=sal_psi,sal_fis=sal_fis,cal_vid=cal_vid,
                viv1=viv1,viv2=viv2,sust_pp=sust_pp,
                sexo=sexo,fn_col=fn_col,fecha=fecha)

def norm_sust(s):
    if pd.isna(s) or str(s).strip()=='0': return None
    s=str(s).strip().lower()
    if any(x in s for x in ['alcohol','cerveza','licor','aguard','alchol','bebida']): return 'Alcohol'
    if any(x in s for x in ['marihu','marjhu','marhuana','cannabis','cannbis']): return 'Cannabis/Marihuana'
    if any(x in s for x in ['pasta base','pasta','papelillo']): return 'Pasta Base'
    if any(x in s for x in ['crack','cristal','piedra','paco']): return 'Crack/Cristal'
    if any(x in s for x in ['cocain','perico','coca ']): return 'Cocaína'
    if any(x in s for x in ['tabaco','cigarr','nicot']): return 'Tabaco/Nicotina'
    if any(x in s for x in ['inhalant','thiner','activo','resistol']): return 'Inhalantes'
    if any(x in s for x in ['sedant','benzod','tranqui']): return 'Sedantes'
    if any(x in s for x in ['opiod','heroina','morfin','fentanil']): return 'Opiáceos'
    if any(x in s for x in ['metanfet','anfetam']): return 'Metanfetamina'
    return 'Otras'

def _es_positivo(valor):
    s=str(valor).strip().lower()
    if s in ('sí','si'): return True
    if s in ('no','no aplica','nunca','nan',''): return False
    n=pd.to_numeric(valor,errors='coerce')
    return not pd.isna(n) and n>0

# ══════════════════════════════════════════════════════════════════════════════
# CARGA DE DATOS
# ══════════════════════════════════════════════════════════════════════════════
def cargar_datos():
    global NOMBRE_SERVICIO, PERIODO, OUTPUT_FILE
    print(f'  Leyendo: {INPUT_FILE}')

    # Detectar país y período desde el archivo Wide
    _pais = _detectar_pais(INPUT_FILE)
    NOMBRE_SERVICIO = _pais if _pais else 'Servicio de Tratamiento'

    _periodo_auto = None
    try:
        _rs = pd.read_excel(INPUT_FILE, sheet_name='Resumen', header=None)
        for _, _row in _rs.iterrows():
            for _v in _row.tolist():
                if 'Período' in str(_v) or 'periodo' in str(_v).lower(): continue
                if '–' in str(_v) or (' ' in str(_v) and any(
                        m in str(_v) for m in ['Enero','Feb','Mar','Abr','May','Jun',
                                               'Jul','Ago','Sep','Oct','Nov','Dic','2024','2025','2026'])):
                    _periodo_auto = str(_v).strip(); break
            if _periodo_auto: break
    except: pass
    PERIODO = _periodo_auto if _periodo_auto else '2025'

    df = pd.read_excel(INPUT_FILE, sheet_name=SHEET_NAME, header=1)
    _col_centro = next((c for c in df.columns if any(x in _norm(c) for x in
                        ['codigo del centro','servicio de tratamiento',
                         'centro/ servicio','codigo centro'])), None)
    if FILTRO_CENTRO and _col_centro:
        n_antes=len(df)
        df=df[df[_col_centro].astype(str).str.strip()==FILTRO_CENTRO].copy().reset_index(drop=True)
        print(f'  ⚑ Filtro: {FILTRO_CENTRO} ({n_antes}→{len(df)} pacientes)')
        _pl=_detectar_pais(INPUT_FILE)
        NOMBRE_SERVICIO=f'{_pl}  —  Centro {FILTRO_CENTRO}' if _pl else f'Centro {FILTRO_CENTRO}'
    N=len(df); DC=detectar_columnas(df.columns.tolist()); R={'N':N,'d':df,'DC':DC}

    # Sexo
    if DC['sexo']:
        sc=df[DC['sexo']].astype(str).str.strip().str.upper()
        nv=int(sc.isin(['H','M']).sum())
        R['n_hombre']=int((sc=='H').sum()); R['n_mujer']=int((sc=='M').sum())
        R['nv_sex']=nv
        R['pct_hombre']=round(R['n_hombre']/nv*100,1) if nv>0 else 0
        R['pct_mujer']=round(R['n_mujer']/nv*100,1) if nv>0 else 0
    else:
        R['n_hombre']=R['n_mujer']=R['nv_sex']=0; R['pct_hombre']=R['pct_mujer']=0

    # Edad
    if DC['fn_col'] and DC['fecha']:
        fn=pd.to_datetime(df[DC['fn_col']],errors='coerce')
        ref=pd.to_datetime(df[DC['fecha']],errors='coerce').fillna(pd.Timestamp.now())
        edad=((ref-fn).dt.days/365.25).round(1); edad=edad[(edad>=10)&(edad<=100)]
        R['nv_edad']=int(edad.notna().sum())
        R['edad_media']=round(float(edad.mean()),1) if R['nv_edad']>0 else 0
        R['edad_sd']=round(float(edad.std()),1) if R['nv_edad']>0 else 0
        R['edad_min']=int(edad.min()) if R['nv_edad']>0 else 0
        R['edad_max']=int(edad.max()) if R['nv_edad']>0 else 0
        bins=[0,17,30,40,50,60,200]; labs=['Menos de 18','18 a 30','31 a 40','41 a 50','51 a 60','61 o más']
        ec=pd.cut(edad,bins=bins,labels=labs)
        R['edad_dist']={l:int((ec==l).sum()) for l in labs}
    else:
        R['edad_media']=R['edad_sd']=0; R['edad_min']=R['edad_max']=0
        R['nv_edad']=0; R['edad_dist']={'Sin datos':0}

    # Sustancia principal
    if DC['sust_pp']:
        sr=df[DC['sust_pp']].apply(norm_sust).dropna()
        R['nv_sust']=len(sr); vc=sr.value_counts(); R['sust_vc']=vc
        R['sust_top1']=vc.index[0] if len(vc)>0 else '—'
        R['sust_top1_pct']=round(vc.iloc[0]/len(sr)*100,1) if len(sr)>0 else 0
    else:
        R['nv_sust']=0; R['sust_vc']=pd.Series(dtype=int)
        R['sust_top1']='—'; R['sust_top1_pct']=0

    # Días consumo sustancia principal
    sust_norm=df[DC['sust_pp']].apply(norm_sust) if DC['sust_pp'] else pd.Series([None]*N)
    dias_princ={}
    for lbl,col in DC['sust_cols']:
        v=pd.to_numeric(df[col],errors='coerce')
        mask=sust_norm.apply(lambda s: isinstance(s,str) and lbl.lower() in s.lower())
        sub=v[mask&(v>0)]
        if len(sub)>=1: dias_princ[lbl]={'prom':round(float(sub.mean()),1),'n':int(len(sub))}
    R['dias_princ']=dias_princ

    # % consumidores
    consumo_pct={}
    for lbl,col in DC['sust_cols']:
        v=pd.to_numeric(df[col],errors='coerce'); n_c=int((v>0).sum())
        if n_c>0: consumo_pct[lbl]={'pct':round(n_c/N*100,1),'n':n_c}
    R['consumo_pct']=consumo_pct

    # Días por sustancia
    dias_sust={}
    for lbl,col in DC['sust_cols']:
        v=pd.to_numeric(df[col],errors='coerce'); sub=v[v>0]
        if len(sub)>=1: dias_sust[lbl]={'prom':round(float(sub.mean()),1),'n':int(len(sub))}
    R['dias_sust']=dias_sust

    # Salud
    salud=[]
    for lbl,col in [('Salud Psicológica',DC['sal_psi']),('Salud Física',DC['sal_fis']),('Calidad de Vida',DC['cal_vid'])]:
        if col:
            v=pd.to_numeric(df[col],errors='coerce')
            salud.append({'label':lbl,'prom':round(float(v.mean()),1),'nv':int(v.notna().sum())})
    R['salud']=salud

    # Vivienda
    def viv(col):
        if not col: return (0,0,0)
        nv_=int(df[col].isin(['Sí','No']).sum()) or N
        n_=int((df[col]=='Sí').sum())
        return n_,round(n_/nv_*100,1),nv_
    R['viv1']=viv(DC['viv1']); R['viv2']=viv(DC['viv2'])

    # Transgresión
    tr_cols=[c for _,c in DC['tr_sn']]
    def has_tr(row):
        for c in tr_cols:
            if _es_positivo(row.get(c,'')): return True
        if DC['vif']:
            v=pd.to_numeric(row.get(DC['vif'],np.nan),errors='coerce')
            return not np.isnan(v) and v>0
        return False
    t=df.apply(lambda r:int(has_tr(r)),axis=1)
    R['n_transgresores']=int(t.sum())
    R['pct_transgresores']=round(R['n_transgresores']/N*100,1) if N>0 else 0
    tipos=[]
    for lbl,col in DC['tr_sn']:
        n=int(df[col].apply(_es_positivo).sum())
        tipos.append({'label':lbl,'n':n,
                      'pct':round(n/R['n_transgresores']*100,1) if R['n_transgresores']>0 else 0})
    if DC['vif']:
        vif_v=pd.to_numeric(df[DC['vif']],errors='coerce'); n_vif=int((vif_v>0).sum())
        tipos.append({'label':'VIF','n':n_vif,
                      'pct':round(n_vif/R['n_transgresores']*100,1) if R['n_transgresores']>0 else 0})
    R['transgtipos']=tipos
    return R

# ══════════════════════════════════════════════════════════════════════════════
# GRÁFICOS
# ══════════════════════════════════════════════════════════════════════════════
def g_sexo(R):
    fig,ax=plt.subplots(figsize=(5,3.5))
    vals=[R['n_hombre'],R['n_mujer']]
    bars=ax.bar(['Hombre','Mujer'],vals,color=[MC_MID,MC_ACCENT],width=0.5,zorder=3)
    for bar,val in zip(bars,vals):
        pct=round(val/R['nv_sex']*100,1) if R['nv_sex']>0 else 0
        ax.text(bar.get_x()+bar.get_width()/2,bar.get_height()+0.5,
                f'{val}\n({pct}%)',ha='center',va='bottom',fontsize=11,fontweight='bold',color='#333')
    ax.set_ylim(0,max(vals)*1.3 if max(vals)>0 else 1)
    ax.set_ylabel('N personas',fontsize=9,color='#595959')
    ax.tick_params(labelsize=10); _ax_style(ax); fig.patch.set_facecolor('white'); fig.tight_layout()
    return fig

def g_edad(R):
    fig,ax=plt.subplots(figsize=(6,4))
    labs=list(R['edad_dist'].keys()); vals=list(R['edad_dist'].values())
    cols=[MC_MID if v==max(vals) else MC_LIGHT for v in vals]
    bars=ax.barh(labs,vals,color=cols,zorder=3)
    for bar,val in zip(bars,vals):
        pct=round(val/R['nv_edad']*100,1) if R['nv_edad']>0 else 0
        ax.text(bar.get_width()+0.2,bar.get_y()+bar.get_height()/2,
                f'{val} ({pct}%)',va='center',fontsize=9,color='#333')
    ax.set_xlim(0,max(vals)*1.5 if max(vals)>0 else 1)
    ax.tick_params(labelsize=9); _ax_style(ax,horiz=True)
    fig.patch.set_facecolor('white'); fig.tight_layout()
    return fig

def g_torta_sust(R):
    vc=R['sust_vc']
    if len(vc)==0:
        fig,ax=plt.subplots(figsize=(6,4)); ax.text(0.5,0.5,'Sin datos',ha='center'); return fig
    labels=list(vc.index); vals=list(vc.values)
    fig,ax=plt.subplots(figsize=(6,4.5))
    wedges,_,autotexts=ax.pie(vals,labels=None,colors=PIE_COLS[:len(vals)],
        autopct=lambda p:f'{p:.1f}%' if p>3 else '',startangle=140,pctdistance=0.72,
        wedgeprops={'edgecolor':'white','linewidth':2.0})
    for at in autotexts: at.set_fontsize(9.5); at.set_color('white'); at.set_fontweight('bold')
    ax.legend(wedges,[f'{l} (n={v})' for l,v in zip(labels,vals)],
              loc='lower center',bbox_to_anchor=(0.5,-0.18),ncol=2,fontsize=8,frameon=False)
    ax.set_aspect('equal'); ax.set_facecolor('white'); fig.patch.set_facecolor('white'); fig.tight_layout()
    return fig

def g_barras_h(datos, ylabel='Promedio días (0–28)', fmt='{v}d\n(n={n})'):
    if not datos: return None
    labs=list(datos.keys()); proms=[datos[l]['prom'] for l in labs]; ns=[datos[l]['n'] for l in labs]
    fig,ax=plt.subplots(figsize=(max(5,len(labs)*1.0),3.8))
    cols=[MC_MID if p==max(proms) else MC_LIGHT for p in proms]
    bars=ax.bar(labs,proms,color=cols,width=0.55,zorder=3)
    for bar,p,n in zip(bars,proms,ns):
        ax.text(bar.get_x()+bar.get_width()/2,bar.get_height()+0.15,
                fmt.format(v=p,n=n),ha='center',va='bottom',fontsize=9,fontweight='bold',color='#333')
    ax.set_ylim(0,max(proms)*1.38); ax.set_ylabel(ylabel,fontsize=9,color='#595959')
    ax.tick_params(labelsize=9); _ax_style(ax); fig.patch.set_facecolor('white'); fig.tight_layout()
    return fig

def g_consumo_pct(R):
    datos=R['consumo_pct']
    if not datos: return None
    labs=list(datos.keys()); vals=[datos[l]['pct'] for l in labs]; ns=[datos[l]['n'] for l in labs]
    fig,ax=plt.subplots(figsize=(max(5,len(labs)*1.0),3.8))
    cols=[MC_MID if v==max(vals) else MC_LIGHT for v in vals]
    bars=ax.bar(labs,vals,color=cols,width=0.55,zorder=3)
    for bar,v,n in zip(bars,vals,ns):
        ax.text(bar.get_x()+bar.get_width()/2,bar.get_height()+0.4,
                f'{v}%\n(n={n})',ha='center',va='bottom',fontsize=9,fontweight='bold',color='#333')
    ax.set_ylim(0,max(vals)*1.4); ax.set_ylabel('% de personas',fontsize=9,color='#595959')
    ax.tick_params(labelsize=9); _ax_style(ax); fig.patch.set_facecolor('white'); fig.tight_layout()
    return fig

def g_donut(R):
    n_si=R['n_transgresores']; n_no=R['N']-n_si
    fig,ax=plt.subplots(figsize=(5,4.5))
    wedges,_,autotexts=ax.pie([n_si,n_no],labels=None,colors=[MC_MID,MC_LIGHT],
        autopct='%1.1f%%',startangle=90,pctdistance=0.75,
        wedgeprops={'edgecolor':'white','linewidth':2.5,'width':0.52},counterclock=False)
    for at in autotexts: at.set_fontsize(13); at.set_fontweight('bold'); at.set_color('white')
    ax.legend(wedges,[f'Cometió transgresión (n={n_si})',f'Sin transgresión (n={n_no})'],
              loc='lower center',bbox_to_anchor=(0.5,-0.08),fontsize=9,frameon=False)
    ax.set_aspect('equal'); ax.set_facecolor('white'); fig.patch.set_facecolor('white'); fig.tight_layout()
    return fig

def g_tipos(R):
    tipos=R['transgtipos']
    labs=[t['label'] for t in tipos]; vals=[t['pct'] for t in tipos]; ns=[t['n'] for t in tipos]
    fig,ax=plt.subplots(figsize=(5,max(3.0,len(labs)*0.6)))
    cols=[MC_MID if v==max(vals) else MC_LIGHT for v in vals]
    bars=ax.barh(labs,vals,color=cols,zorder=3)
    for bar,v,n in zip(bars,vals,ns):
        ax.text(bar.get_width()+0.5,bar.get_y()+bar.get_height()/2,
                f'{v}% (n={n})',va='center',fontsize=9,color='#333')
    ax.set_xlim(0,max(vals)*1.6 if vals else 1)
    ax.set_xlabel(f'% sobre {R["n_transgresores"]} transgresores',fontsize=8.5,color='#595959')
    ax.tick_params(labelsize=9.5); _ax_style(ax,horiz=True)
    fig.patch.set_facecolor('white'); fig.tight_layout()
    return fig

def g_salud(R):
    labels=[s['label'] for s in R['salud']]; proms=[s['prom'] for s in R['salud']]
    fig,ax=plt.subplots(figsize=(5,3.2))
    bars=ax.barh(labels,proms,color=[MC_MID,MC_LIGHT,MC_ACCENT],zorder=3)
    for bar,p in zip(bars,proms):
        ax.text(bar.get_width()+0.1,bar.get_y()+bar.get_height()/2,
                f'{p}/20',va='center',fontsize=10,fontweight='bold',color='#333')
    ax.set_xlim(0,24); ax.axvline(x=10,color='#BFBFBF',linestyle='--',linewidth=1.0)
    ax.tick_params(labelsize=10); _ax_style(ax,horiz=True)
    fig.patch.set_facecolor('white'); fig.tight_layout()
    return fig

def g_vivienda(R):
    fig,ax=plt.subplots(figsize=(5,3.5))
    cats=['Lugar estable','Condiciones básicas']
    n_si=[R['viv1'][0],R['viv2'][0]]; n_no=[R['viv1'][2]-R['viv1'][0],R['viv2'][2]-R['viv2'][0]]
    nv=[R['viv1'][2],R['viv2'][2]]
    x,w2=np.arange(len(cats)),0.35
    b1=ax.bar(x-w2/2,n_si,w2,label='Sí',color=MC_MID,zorder=3)
    b2=ax.bar(x+w2/2,n_no,w2,label='No',color=MC_GRAY,zorder=3)
    for bar,n,nv_ in zip(list(b1)+list(b2),n_si+n_no,nv+nv):
        pct=round(n/nv_*100,1) if nv_>0 else 0
        ax.text(bar.get_x()+bar.get_width()/2,bar.get_height()+0.3,
                f'{n}\n({pct}%)',ha='center',va='bottom',fontsize=9,color='#333')
    ax.set_xticks(x); ax.set_xticklabels(cats,fontsize=10)
    ax.set_ylim(0,max(n_si+n_no)*1.45 if max(n_si+n_no)>0 else 1)
    ax.set_ylabel('N personas',fontsize=9,color='#595959')
    ax.legend(fontsize=9,frameon=False); _ax_style(ax)
    fig.patch.set_facecolor('white'); fig.tight_layout()
    return fig

# ══════════════════════════════════════════════════════════════════════════════
# CONSTRUCCIÓN DEL WORD
# ══════════════════════════════════════════════════════════════════════════════
def build_word(R):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
        sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)

    # ── Portada ───────────────────────────────────────────────────────────────
    from docx.oxml.ns import qn as _qn
    tbl=doc.add_table(rows=1,cols=1); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    c=tbl.rows[0].cells[0]; set_cell_bg(c,C_NAVY)
    for txt,sz,bold in [
        ('INFORME DE CARACTERIZACIÓN',18,True),
        ('Monitoreo de Resultados de Tratamiento — Instrumento TOP',11,False),
        (NOMBRE_SERVICIO.upper(),14,True),
        (PERIODO,10,False),
    ]:
        p=c.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=p.add_run(txt); run.font.name='Calibri'; run.font.size=Pt(sz)
        run.font.bold=bold; run.font.color.rgb=WHITE
    doc.add_paragraph()

    # KPIs
    add_kpi_row(doc,[
        (R['N'],'Personas ingresaron'),
        (f'{R["pct_hombre"]}%','Son hombres'),
        (R['edad_media'],'Edad promedio'),
        (R['sust_top1'],f'Sust. principal ({R["sust_top1_pct"]}%)'),
    ])

    # Presentación
    add_section_header(doc,'','Presentación')
    add_body(doc,
        f'El presente informe describe el perfil de las personas que ingresan a tratamiento '
        f'por consumo de sustancias en {NOMBRE_SERVICIO}, durante el período {PERIODO}, '
        f'a través del instrumento TOP. Durante este período ingresaron {R["N"]} personas; '
        f'la sustancia de mayor problema al ingreso fue {R["sust_top1"]} ({R["sust_top1_pct"]}%). '
        f'El {R["pct_hombre"]}% son hombres y el {R["pct_mujer"]}% son mujeres.')
    doc.add_paragraph()

    # ── Sección 1: Antecedentes generales ─────────────────────────────────────
    # Salto de página explícito: garantiza que el banner + gráfico 1.1 arranquen juntos
    p_break = doc.add_paragraph()
    run_break = p_break.add_run()
    run_break.add_break(__import__('docx.enum.text', fromlist=['WD_BREAK']).WD_BREAK.PAGE)

    add_section_header(doc,'1','Antecedentes Generales')

    add_subsection(doc,'1.1. Distribución de Personas según Sexo')
    buf,w=fig_to_img(g_sexo(R),11)
    add_picture_kwnext(doc,buf,w)
    add_body(doc,
        f'Del total de {R["N"]} personas, {R["n_hombre"]} ({R["pct_hombre"]}%) son hombres '
        f'y {R["n_mujer"]} ({R["pct_mujer"]}%) son mujeres.')
    add_note(doc,f'N válido: {R["nv_sex"]} personas.')
    doc.add_paragraph()

    add_subsection(doc,'1.2. Distribución de Personas según Edad')
    buf,w=fig_to_img(g_edad(R),13)
    add_picture_kwnext(doc,buf,w)
    rm=max(R['edad_dist'],key=R['edad_dist'].get) if R['edad_dist'] else '—'
    nm=R['edad_dist'].get(rm,0); pm=round(nm/R['nv_edad']*100,1) if R['nv_edad']>0 else 0
    add_body(doc,
        f'El promedio de edad es {R["edad_media"]} años (DE={R["edad_sd"]}; rango {R["edad_min"]}–{R["edad_max"]}). '
        f'El rango más frecuente es {rm}, con {nm} personas ({pm}%).')
    add_note(doc,f'N válido: {R["nv_edad"]} personas.')
    doc.add_paragraph()

    # ── Sección 2: Consumo de sustancias ──────────────────────────────────────
    add_section_header(doc,'2','Consumo de Sustancias')

    add_subsection(doc,'2.1. Sustancia Principal al Ingreso')
    buf,w=fig_to_img(g_torta_sust(R),13)
    add_picture_kwnext(doc,buf,w)
    vc=R['sust_vc']
    seg2=vc.index[1] if len(vc)>1 else '—'; pct_s2=round(vc.iloc[1]/R['nv_sust']*100,1) if len(vc)>1 and R['nv_sust']>0 else 0
    add_body(doc,
        f'La sustancia más frecuente es {R["sust_top1"]} ({R["sust_top1_pct"]}%), '
        f'seguida por {seg2} ({pct_s2}%).')
    add_note(doc,f'N válido: {R["nv_sust"]} personas.')
    doc.add_paragraph()

    if R['dias_princ']:
        add_subsection(doc,'2.2. Promedio de Días de Consumo por Sustancia Principal')
        fig=g_barras_h(R['dias_princ'])
        if fig:
            buf,w=fig_to_img(fig,13)
            add_picture_kwnext(doc,buf,w)
            dp_max=max(R['dias_princ'],key=lambda k:R['dias_princ'][k]['prom'])
            add_body(doc,
                f'{dp_max} presenta el mayor promedio: {R["dias_princ"][dp_max]["prom"]} días '
                f'(n={R["dias_princ"][dp_max]["n"]}). '
                f'Promedio en últimas 4 semanas, calculado sobre quienes declararon esa sustancia como principal.')
        doc.add_paragraph()

    if R['consumo_pct']:
        add_subsection(doc,'2.3. Consumo de Sustancias — % de Personas')
        fig=g_consumo_pct(R)
        if fig:
            buf,w=fig_to_img(fig,13)
            add_picture_kwnext(doc,buf,w)
            cp=R['consumo_pct']; sk=max(cp,key=lambda k:cp[k]['pct'])
            add_body(doc,
                f'Los porcentajes pueden sumar más del 100% ya que una persona puede consumir varias sustancias. '
                f'{sk} es la más prevalente: {cp[sk]["pct"]}% ({cp[sk]["n"]} personas).')
            add_note(doc,f'N total: {R["N"]} personas.')
        doc.add_paragraph()

    if R['dias_sust']:
        add_subsection(doc,'2.4. Promedio de Días de Consumo por Sustancia')
        fig=g_barras_h(R['dias_sust'])
        if fig:
            buf,w=fig_to_img(fig,13)
            add_picture_kwnext(doc,buf,w)
            ds=R['dias_sust']; dk=max(ds,key=lambda k:ds[k]['prom'])
            add_body(doc,
                f'{dk} tiene el mayor promedio: {ds[dk]["prom"]} días (n={ds[dk]["n"]}). '
                f'Promedio calculado solo entre consumidores (días > 0).')
        doc.add_paragraph()

    # ── Sección 3: Transgresión ────────────────────────────────────────────────
    p_break3 = doc.add_paragraph()
    p_break3.add_run().add_break(__import__('docx.enum.text', fromlist=['WD_BREAK']).WD_BREAK.PAGE)
    add_section_header(doc,'3','Transgresión a la Norma Social')

    add_subsection(doc,'3.1. Transgresión a la Norma Social')
    buf,w=fig_to_img(g_donut(R),12)
    add_picture_kwnext(doc,buf,w)
    n_no_tr=R['N']-R['n_transgresores']; pct_no=round(n_no_tr/R['N']*100,1) if R['N']>0 else 0
    add_body(doc,
        f'{R["n_transgresores"]} personas ({R["pct_transgresores"]}%) declararon haber cometido '
        f'algún tipo de transgresión en el mes previo al ingreso. '
        f'Las {n_no_tr} personas restantes ({pct_no}%) no reportaron ningún incidente.')
    add_note(doc,f'N total: {R["N"]} personas.')
    doc.add_paragraph()

    if R['transgtipos']:
        add_subsection(doc,'3.2. Distribución por Tipo de Transgresión')
        buf,w=fig_to_img(g_tipos(R),13)
        add_picture_kwnext(doc,buf,w)
        tm=max(R['transgtipos'],key=lambda t:t['n'])
        add_body(doc,
            f'El tipo más frecuente es {tm["label"]} ({tm["pct"]}%, n={tm["n"]}). '
            f'Los porcentajes no suman 100% porque una persona puede haber cometido más de un tipo.')
        add_note(doc,f'N base: {R["n_transgresores"]} personas con al menos una transgresión.')
        doc.add_paragraph()

    # ── Sección 4: Salud y vivienda ────────────────────────────────────────────
    add_section_header(doc,'4','Salud, Calidad de Vida y Vivienda')

    if R['salud']:
        add_subsection(doc,'4.1. Autopercepción del Estado de Salud')
        buf,w=fig_to_img(g_salud(R),13)
        add_picture_kwnext(doc,buf,w)
        mejor=max(R['salud'],key=lambda s:s['prom'])
        add_body(doc,
            f'Puntajes promedio de autopercepción en escala 0–20 (0=muy mal, 20=excelente). '
            f'La dimensión mejor evaluada es {mejor["label"]} ({mejor["prom"]}/20). '
            f'Puntajes bajo 10 indican percepción deficiente.')
        add_note(doc,f'N válido: {R["salud"][0]["nv"]} personas.')
        doc.add_paragraph()

    add_subsection(doc,'4.2. Condiciones de Vivienda al Ingreso')
    buf,w=fig_to_img(g_vivienda(R),13)
    add_picture_kwnext(doc,buf,w)
    n1,p1,nv1=R['viv1']; n2,p2,nv2=R['viv2']
    add_body(doc,
        f'El {p1}% de las personas ({n1} de {nv1}) declaran tener un lugar estable donde vivir. '
        f'El {p2}% ({n2} de {nv2}) habita en una vivienda que cumple condiciones básicas.')
    doc.add_paragraph()

    # ── Pie ───────────────────────────────────────────────────────────────────
    p=doc.add_paragraph()
    run=p.add_run(f'Informe generado automáticamente · TOP · {NOMBRE_SERVICIO} · {PERIODO}')
    run.font.size=Pt(8); run.font.italic=True; run.font.color.rgb=GRAY
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUTPUT_FILE)
    print(f'  ✓ Word generado: {OUTPUT_FILE}')

# ══════════════════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    print('='*60)
    print('  SCRIPT_TOP_Universal_Word_Caracterizacion  —  Iniciando...')
    print('='*60)
    R = cargar_datos()
    print(f'  N={R["N"]} | {R["sust_top1"]} {R["sust_top1_pct"]}% | Transgr.:{R["n_transgresores"]}')
    build_word(R)
    print(f'\n{"="*60}')
    print(f'  ✅  LISTO  →  {OUTPUT_FILE}')
    print(f'{"="*60}')
