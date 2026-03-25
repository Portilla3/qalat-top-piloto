"""
╔══════════════════════════════════════════════════════════════════════════════╗
║   SCRIPT_TOP_Universal_Word_Seguimiento.py  —  v1.0                       ║
║   Genera informe Word comparativo TOP1 vs TOP2                             ║
║   11 gráficos · Compatible con cualquier país TOP                          ║
╠══════════════════════════════════════════════════════════════════════════════╣
║  CÓMO USAR:                                                                 ║
║  1. Abre un chat nuevo con Claude                                           ║
║  2. Sube DOS archivos:                                                      ║
║       • Este script                                                         ║
║       • La base Wide (generada por SCRIPT_TOP_Universal_Wide)              ║
║  3. Escribe: "Ejecuta el script Word Seguimiento TOP con esta base"        ║
╚══════════════════════════════════════════════════════════════════════════════╝
"""

import glob, os, unicodedata, io, warnings
import pandas as pd
import numpy as np
import matplotlib; matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
warnings.filterwarnings('ignore')

NAVY=RGBColor(0x1F,0x38,0x64); BLUE=RGBColor(0x2E,0x75,0xB6)
ACCENT=RGBColor(0x00,0xB0,0xF0); GRAY=RGBColor(0x59,0x59,0x59)
WHITE=RGBColor(0xFF,0xFF,0xFF)
C_NAVY='1F3864'; C_MID='2E75B6'; C_LIGHT='EEF4FB'
MC_MID='#2E75B6'; MC_LIGHT='#BDD7EE'; MC_ACCENT='#00B0F0'
C_T1='#2E75B6'; C_T2='#00B0F0'
C_ABS2='#1F3864'; C_DIS2='#2E75B6'; C_SC2='#9DC3E6'; C_EMP2='#BDD7EE'
PIE_COLS=['#2E75B6','#1F3864','#4472C4','#9DC3E6','#00B0F0','#538135','#BFBFBF','#C00000','#ED7D31']

_PAISES = {
    'republica_dominicana':'República Dominicana','dominicana':'República Dominicana',
    'honduras':'Honduras','panama':'Panamá','panam':'Panamá',
    'el_salvador':'El Salvador','salvador':'El Salvador',
    'mexico':'México','mexic':'México','ecuador':'Ecuador',
    'peru':'Perú','argentina':'Argentina','colombia':'Colombia',
    'chile':'Chile','bolivia':'Bolivia','paraguay':'Paraguay',
    'uruguay':'Uruguay','venezuela':'Venezuela','guatemala':'Guatemala',
    'costa_rica':'Costa Rica','costarica':'Costa Rica','nicaragua':'Nicaragua',
}

def _norm(s): return unicodedata.normalize('NFD',str(s).lower()).encode('ascii','ignore').decode()
def _extraer_pais(fn):
    f=_norm(str(fn).replace('.','_'))
    for k,v in _PAISES.items():
        if k in f: return v
    return None

def _detectar_pais(wide_file):
    try:
        rs=pd.read_excel(wide_file,sheet_name='Resumen',header=None)
        for _,row in rs.iterrows():
            for v in row.tolist():
                p=_extraer_pais(str(v))
                if p: return p
    except: pass
    return _extraer_pais(os.path.basename(wide_file))

def auto_archivo_wide():
    candidatos=(
        glob.glob('/mnt/user-data/uploads/*Wide*.xlsx')+
        glob.glob('/mnt/user-data/uploads/*wide*.xlsx')+
        glob.glob('/mnt/user-data/uploads/TOP_Base*.xlsx')+
        glob.glob('/home/claude/TOP_Base_Wide.xlsx'))
    if not candidatos: raise FileNotFoundError('No se encontró la base Wide TOP.')
    print(f'  → Base Wide: {os.path.basename(candidatos[0])}')
    return candidatos[0]

# ══════════════════════════════════════════════════════════════════════════════
# CONFIGURACIÓN — inyectada por runner.py (no ejecutar a nivel módulo)
# ══════════════════════════════════════════════════════════════════════════════
INPUT_FILE    = None   # runner inyecta la ruta real
SHEET_NAME    = 'Base Wide'
OUTPUT_FILE   = None   # runner inyecta la ruta real
FILTRO_CENTRO = None   # runner inyecta el filtro si aplica
NOMBRE_SERVICIO = 'Servicio de Tratamiento'
PERIODO         = ''

# ── Helpers Word ──────────────────────────────────────────────────────────────
def set_cell_bg(cell,hex_color):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear')
    shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hex_color); tcPr.append(shd)

def _keep_with_next(p):
    """Evita que el párrafo quede separado del siguiente elemento."""
    pPr = p._p.get_or_add_pPr()
    pPr.append(OxmlElement('w:keepNext'))
    pPr.append(OxmlElement('w:keepLines'))

def add_section_header(doc,num,title):
    tbl=doc.add_table(rows=1,cols=1); tbl.alignment=WD_TABLE_ALIGNMENT.LEFT
    c=tbl.rows[0].cells[0]; set_cell_bg(c,C_MID)
    p=c.paragraphs[0]; p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
    txt=f'  {num}. {title.upper()}' if num else f'  {title.upper()}'
    run=p.add_run(txt); run.font.name='Calibri'; run.font.size=Pt(11)
    run.font.bold=True; run.font.color.rgb=WHITE
    sp=doc.add_paragraph(); sp.paragraph_format.space_after=Pt(0)
    _keep_with_next(sp)

def add_subsection(doc,title):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(2)
    run=p.add_run(title); run.bold=True; run.font.name='Calibri'; run.font.size=Pt(10.5)
    run.font.color.rgb=NAVY; _keep_with_next(p)

def add_picture_kwnext(doc,buf,width):
    """Inserta imagen centrada con keepNext para que el texto no se separe."""
    doc.add_picture(buf,width=width)
    p=doc.paragraphs[-1]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _keep_with_next(p)

def add_body(doc,text,italic=False):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3)
    run=p.add_run(text); run.font.name='Calibri'; run.font.size=Pt(10)
    run.font.italic=italic; run.font.color.rgb=RGBColor(0x33,0x33,0x33); return p

def add_note(doc,text):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
    run=p.add_run(text); run.font.name='Calibri'; run.font.size=Pt(8.5)
    run.font.italic=True; run.font.color.rgb=GRAY

def add_kpi_row(doc,kpis):
    tbl=doc.add_table(rows=2,cols=len(kpis)); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,(val,lbl) in enumerate(kpis):
        cv=tbl.rows[0].cells[j]; set_cell_bg(cv,C_LIGHT)
        p=cv.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=p.add_run(str(val)); run.font.name='Calibri'; run.font.size=Pt(20)
        run.font.bold=True; run.font.color.rgb=BLUE
        cl=tbl.rows[1].cells[j]; set_cell_bg(cl,C_LIGHT)
        p2=cl.paragraphs[0]; p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run2=p2.add_run(lbl); run2.font.name='Calibri'; run2.font.size=Pt(8.5)
        run2.font.color.rgb=GRAY
    doc.add_paragraph()

def fig_to_img(fig,width_cm=13):
    buf=io.BytesIO()
    fig.savefig(buf,format='png',dpi=150,bbox_inches='tight',facecolor='white')
    buf.seek(0); plt.close(fig); return buf,Cm(width_cm)

def _ax_style(ax,horiz=False):
    (ax.xaxis if horiz else ax.yaxis).grid(True,color='#E2E8F0',linewidth=0.6,zorder=0)
    ax.set_axisbelow(True)
    for sp in ['top','right']: ax.spines[sp].set_visible(False)
    ax.spines['left'].set_color('#D0D0D0'); ax.spines['bottom'].set_color('#D0D0D0')
    ax.set_facecolor('white')

def _es_positivo(valor):
    s=str(valor).strip().lower()
    if s in ('sí','si'): return True
    if s in ('no','no aplica','nunca','nan',''): return False
    n=pd.to_numeric(valor,errors='coerce'); return not pd.isna(n) and n>0

def flecha(v1,v2,mejor_si_sube=True):
    if v1==v2: return 'Sin cambio'
    mejoro=(v2>v1)==mejor_si_sube
    return f'↑ Mejoró ({v1}→{v2})' if mejoro else f'↓ Empeoró ({v1}→{v2})'

# ── Detección de columnas ─────────────────────────────────────────────────────
def detectar_columnas(cols):
    col_set=set(cols)
    def par(c1):
        c2=c1.replace('_TOP1','_TOP2'); return (c1,c2 if c2 in col_set else None)
    sust_cols=[]
    for c in cols:
        if c.endswith('_TOP1') and 'Total (0-28)' in c:
            base=c.replace('_TOP1','')
            if base.startswith('1)'):
                partes=base.split('>>')
                if len(partes)>=3:
                    nombre=partes[-2].strip().split('(')[0].strip()
                    c1,c2=par(c); sust_cols.append((nombre,c1,c2))
    tr_sn=[]
    for c in cols:
        if c.endswith('_TOP1') and c.replace('_TOP1','').startswith('3)') and '>>' in c:
            nombre=c.replace('_TOP1','').split('>>')[-1].strip()
            c1,c2=par(c); tr_sn.append((nombre,c1,c2))
    def safe(p): return p if (p[0] and p[0] in col_set) else (None,None)
    vif=safe(par(next((c for c in cols if c.endswith('_TOP1') and '4)' in c
                 and 'Violencia Intrafamiliar' in c and 'Total (0-28)' in c),'') or ''))
    sal_psi=safe(par(next((c for c in cols if c.endswith('_TOP1') and c.replace('_TOP1','').startswith('6)')),'') or ''))
    sal_fis=safe(par(next((c for c in cols if c.endswith('_TOP1') and c.replace('_TOP1','').startswith('8)')),'') or ''))
    cal_vid=safe(par(next((c for c in cols if c.endswith('_TOP1') and c.replace('_TOP1','').startswith('10)')),'') or ''))
    viv1=safe(par(next((c for c in cols if c.endswith('_TOP1') and '9)' in c and 'estable' in c.lower()),'') or ''))
    viv2=safe(par(next((c for c in cols if c.endswith('_TOP1') and '9)' in c and 'básicas' in c.lower()),'') or ''))
    sust_pp=safe(par(next((c for c in cols if c.endswith('_TOP1') and c.replace('_TOP1','').startswith('2)')
                      and 'sustancia principal' in c.lower()),'') or ''))
    sexo=next((c for c in cols if c.endswith('_TOP1') and 'sexo' in c.lower()),None)
    fn_col=next((c for c in cols if c.endswith('_TOP1') and 'nacimiento' in c.lower()),None)
    fecha=next((c for c in cols if c.endswith('_TOP1') and 'fecha entrevista' in c.lower()),None)
    return dict(sust_cols=sust_cols,tr_sn=tr_sn,vif=vif,sal_psi=sal_psi,sal_fis=sal_fis,
                cal_vid=cal_vid,viv1=viv1,viv2=viv2,sust_pp=sust_pp,sexo=sexo,fn_col=fn_col,fecha=fecha)

def norm_sust(s):
    if pd.isna(s) or str(s).strip()=='0': return None
    s=str(s).strip().lower()
    if any(x in s for x in ['alcohol','cerveza','licor','aguard','alchol','bebida']): return 'Alcohol'
    if any(x in s for x in ['marihu','marjhu','marhuana','cannabis','cannbis']): return 'Cannabis/Marihuana'
    if any(x in s for x in ['pasta base','pasta','papelillo']): return 'Pasta Base'
    if any(x in s for x in ['crack','cristal','piedra','paco']): return 'Crack/Cristal'
    if any(x in s for x in ['cocain','perico','coca ']): return 'Cocaína'
    if any(x in s for x in ['tabaco','cigarr','nicot']): return 'Tabaco'
    if any(x in s for x in ['sedant','benzod','tranqui']): return 'Sedantes'
    if any(x in s for x in ['opiod','heroina','morfin','fentanil']): return 'Opiáceos'
    if any(x in s for x in ['metanfet','anfetam']): return 'Metanfetamina'
    return 'Otras'

# ── Carga de datos ────────────────────────────────────────────────────────────
def cargar_datos():
    global NOMBRE_SERVICIO, PERIODO, OUTPUT_FILE
    print(f'  Leyendo: {INPUT_FILE}')

    # Detectar país y período desde el archivo Wide
    _pais = _detectar_pais(INPUT_FILE)
    NOMBRE_SERVICIO = _pais if _pais else 'Servicio de Tratamiento'

    _periodo_auto = None
    try:
        _rs=pd.read_excel(INPUT_FILE,sheet_name='Resumen',header=None)
        for _,_row in _rs.iterrows():
            for _v in _row.tolist():
                if 'Período' in str(_v) or 'periodo' in str(_v).lower(): continue
                if '–' in str(_v) or (' ' in str(_v) and any(
                        m in str(_v) for m in ['Enero','Feb','Mar','Abr','May','Jun',
                                               'Jul','Ago','Sep','Oct','Nov','Dic','2024','2025','2026'])):
                    _periodo_auto=str(_v).strip(); break
            if _periodo_auto: break
    except: pass
    PERIODO = _periodo_auto if _periodo_auto else '2025'

    df=pd.read_excel(INPUT_FILE,sheet_name=SHEET_NAME,header=1)
    _col_centro=next((c for c in df.columns if any(x in _norm(c) for x in
                      ['codigo del centro','servicio de tratamiento',
                       'centro/ servicio','codigo centro'])),None)
    if FILTRO_CENTRO and _col_centro:
        n_antes=len(df)
        df=df[df[_col_centro].astype(str).str.strip()==FILTRO_CENTRO].copy().reset_index(drop=True)
        print(f'  ⚑ Filtro: {FILTRO_CENTRO} ({n_antes}→{len(df)} pacientes)')
        _pl=_detectar_pais(INPUT_FILE)
        NOMBRE_SERVICIO=f'{_pl}  —  Centro {FILTRO_CENTRO}' if _pl else f'Centro {FILTRO_CENTRO}'
    N_total=len(df)
    seg=df[df['Tiene_TOP2']=='Sí'].copy().reset_index(drop=True)
    N_seg=len(seg)
    print(f'  Total: {N_total} | Con TOP2: {N_seg} ({round(N_seg/N_total*100,1) if N_total>0 else 0}%)')

    _fc1=next((c for c in seg.columns if 'fecha entrevista' in c.lower() and c.endswith('_TOP1')),None)
    _fc2=next((c for c in seg.columns if 'fecha entrevista' in c.lower() and c.endswith('_TOP2')),None)
    seg_tiempo={'mediana':None,'min':None,'max':None,'n':0,'n_total':N_seg}
    if _fc1 and _fc2:
        _d1=pd.to_datetime(seg[_fc1],errors='coerce'); _d2=pd.to_datetime(seg[_fc2],errors='coerce')
        _dias=(_d2-_d1).dt.days; _dias_ok=_dias[(_dias>=0)&(_dias<=730)].dropna()
        if len(_dias_ok)>0:
            _m=_dias_ok/30.44
            seg_tiempo={'mediana':round(float(_m.median()),1),'min':round(float(_m.min()),1),
                        'max':round(float(_m.max()),1),'n':len(_dias_ok),'n_total':int(_dias.notna().sum())}

    DC=detectar_columnas(seg.columns.tolist())
    R={'N_total':N_total,'N_seg':N_seg,'DC':DC,'seg_tiempo':seg_tiempo}

    # Sexo
    if DC['sexo']:
        sc=seg[DC['sexo']].astype(str).str.strip().str.upper()
        nv=int(sc.isin(['H','M']).sum())
        R['n_hombre']=int((sc=='H').sum()); R['n_mujer']=int((sc=='M').sum()); R['nv_sex']=nv
        R['pct_hombre']=round(R['n_hombre']/nv*100,1) if nv>0 else 0
        R['pct_mujer']=round(R['n_mujer']/nv*100,1) if nv>0 else 0
    else:
        R['n_hombre']=R['n_mujer']=R['nv_sex']=0; R['pct_hombre']=R['pct_mujer']=0

    # Edad
    if DC['fn_col'] and DC['fecha']:
        fn=pd.to_datetime(seg[DC['fn_col']],errors='coerce')
        ref=pd.to_datetime(seg[DC['fecha']],errors='coerce').fillna(pd.Timestamp.now())
        edad=((ref-fn).dt.days/365.25).round(1); edad=edad[(edad>=10)&(edad<=100)]
        R['edad_media']=round(float(edad.mean()),1) if len(edad)>0 else 0
        R['edad_sd']=round(float(edad.std()),1) if len(edad)>0 else 0
        R['nv_edad']=int(edad.notna().sum()); R['edad_min']=int(edad.min()) if len(edad)>0 else 0
        R['edad_max']=int(edad.max()) if len(edad)>0 else 0
        bins=[0,17,30,40,50,60,200]; labs=['Menos de 18','18 a 30','31 a 40','41 a 50','51 a 60','61 o más']
        ec=pd.cut(edad,bins=bins,labels=labs); R['edad_dist']={l:int((ec==l).sum()) for l in labs}
    else:
        R['edad_media']=R['edad_sd']=0; R['edad_min']=R['edad_max']=0
        R['nv_edad']=0; R['edad_dist']={'Sin datos':0}

    # Sustancia principal TOP1 vs TOP2
    c1_sp,c2_sp=DC['sust_pp']
    if c1_sp:
        sr1=seg[c1_sp].apply(norm_sust)
        sr2=seg[c2_sp].apply(norm_sust) if c2_sp else pd.Series([None]*N_seg)
        R['nv_sust1']=int(sr1.notna().sum()); R['nv_sust2']=int(sr2.notna().sum())
        cats=['Alcohol','Cannabis/Marihuana','Pasta Base','Cocaína','Crack/Cristal',
              'Tabaco','Sedantes','Opiáceos','Metanfetamina','Otras']
        sust_comp=[]
        for cat in cats:
            n1=int((sr1==cat).sum()); n2=int((sr2==cat).sum())
            if n1>0 or n2>0:
                sust_comp.append({'label':cat,'n1':n1,'n2':n2,
                    'p1':round(n1/R['nv_sust1']*100,1) if R['nv_sust1']>0 else 0,
                    'p2':round(n2/R['nv_sust2']*100,1) if R['nv_sust2']>0 else 0})
        R['sust_comp']=sust_comp
        top1=max(sust_comp,key=lambda x:x['n1']) if sust_comp else {'label':'—','p1':0}
        R['sust_top1']=top1['label']; R['sust_top1_pct']=top1['p1']
    else:
        R['nv_sust1']=R['nv_sust2']=0; R['sust_comp']=[]; R['sust_top1']='—'; R['sust_top1_pct']=0

    # Días consumo TOP1 vs TOP2
    dias_comp=[]
    for lbl,c1,c2 in DC['sust_cols']:
        v1=pd.to_numeric(seg[c1],errors='coerce')
        v2=pd.to_numeric(seg[c2],errors='coerce') if c2 else pd.Series([np.nan]*N_seg)
        m1=round(float(v1.mean()),1) if v1.notna().sum()>0 else 0
        m2=round(float(v2.mean()),1) if (c2 and v2.notna().sum()>0) else 0
        if m1>0 or m2>0:
            dias_comp.append({'label':lbl,'m1':m1,'m2':m2,
                              'nv1':int(v1.notna().sum()),'nv2':int(v2.notna().sum()) if c2 else 0})
    R['dias_comp']=dias_comp

    # % consumidores TOP1 vs TOP2
    cons_pct=[]
    for lbl,c1,c2 in DC['sust_cols']:
        v1=pd.to_numeric(seg[c1],errors='coerce')
        v2=pd.to_numeric(seg[c2],errors='coerce') if c2 else pd.Series([0]*N_seg)
        n1=int((v1>0).sum()); n2=int((v2>0).sum()) if c2 else 0
        if n1>0 or n2>0:
            cons_pct.append({'label':lbl,'n1':n1,'n2':n2,
                'p1':round(n1/N_seg*100,1),'p2':round(n2/N_seg*100,1) if c2 else 0})
    R['cons_pct']=cons_pct

    # Cambio en consumo
    cambio=[]
    for lbl,c1,c2 in DC['sust_cols']:
        if not c2: continue
        v1=pd.to_numeric(seg[c1],errors='coerce').fillna(0)
        v2=pd.to_numeric(seg[c2],errors='coerce').fillna(0)
        mask=v1>0; n_cons=int(mask.sum())
        if n_cons<2: continue
        s1=v1[mask]; s2=v2[mask]
        n_abs=int((s2==0).sum()); n_dis=int(((s2>0)&(s2<s1)).sum())
        n_sc=int((s2==s1).sum()); n_emp=int((s2>s1).sum())
        pct=lambda n: round(n/n_cons*100,1) if n_cons>0 else 0
        cambio.append({'label':lbl,'n_cons':n_cons,
            'pct_abs':pct(n_abs),'pct_dis':pct(n_dis),'pct_sc':pct(n_sc),'pct_emp':pct(n_emp)})
    R['cambio']=cambio

    # Salud TOP1 vs TOP2
    salud=[]
    for lbl,(c1,c2) in [('Salud Psicológica',DC['sal_psi']),('Salud Física',DC['sal_fis']),('Calidad de Vida',DC['cal_vid'])]:
        if not c1: continue
        v1=pd.to_numeric(seg[c1],errors='coerce')
        v2=pd.to_numeric(seg[c2],errors='coerce') if c2 else pd.Series([np.nan]*N_seg)
        salud.append({'label':lbl,
            'm1':round(float(v1.mean()),1),'m2':round(float(v2.mean()),1) if c2 else 0,
            'nv1':int(v1.notna().sum()),'nv2':int(v2.notna().sum()) if c2 else 0})
    R['salud']=salud

    # Vivienda TOP1 vs TOP2
    def viv(col,df_):
        if not col or col not in df_.columns: return (0,0,N_seg)
        nv=int(df_[col].isin(['Sí','No']).sum()) or N_seg
        n=int((df_[col]=='Sí').sum()); return n,round(n/nv*100,1),nv
    c_viv1_1,c_viv1_2=DC['viv1']; c_viv2_1,c_viv2_2=DC['viv2']
    R['viv1_t1']=viv(c_viv1_1,seg); R['viv1_t2']=viv(c_viv1_2,seg)
    R['viv2_t1']=viv(c_viv2_1,seg); R['viv2_t2']=viv(c_viv2_2,seg)

    # Transgresión TOP1 vs TOP2
    tr_cols1=[c1 for _,c1,_ in DC['tr_sn']]; tr_cols2=[c2 for _,_,c2 in DC['tr_sn']]
    vif_c1,vif_c2=DC['vif']
    def has_tr(row,sn_cols,vif_col):
        for c in sn_cols:
            if c and _es_positivo(row.get(c,'')): return True
        if vif_col:
            v=pd.to_numeric(row.get(vif_col,np.nan),errors='coerce')
            return not np.isnan(v) and v>0
        return False
    tr1=seg.apply(lambda r:int(has_tr(r,tr_cols1,vif_c1)),axis=1)
    tr2=seg.apply(lambda r:int(has_tr(r,tr_cols2,vif_c2)),axis=1)
    R['n_tr1']=int(tr1.sum()); R['n_tr2']=int(tr2.sum())
    R['pct_tr1']=round(R['n_tr1']/N_seg*100,1) if N_seg>0 else 0
    R['pct_tr2']=round(R['n_tr2']/N_seg*100,1) if N_seg>0 else 0
    tipos=[]
    for lbl,c1,c2 in DC['tr_sn']:
        n1=int(seg[c1].apply(_es_positivo).sum()) if c1 else 0
        n2=int(seg[c2].apply(_es_positivo).sum()) if c2 else 0
        tipos.append({'label':lbl,'n1':n1,'n2':n2,
            'p1':round(n1/N_seg*100,1),'p2':round(n2/N_seg*100,1)})
    if vif_c1:
        vif1_v=pd.to_numeric(seg[vif_c1],errors='coerce'); n_v1=int((vif1_v>0).sum())
        vif2_v=pd.to_numeric(seg[vif_c2],errors='coerce') if vif_c2 else pd.Series([np.nan]*N_seg)
        n_v2=int((vif2_v>0).sum())
        tipos.append({'label':'VIF','n1':n_v1,'n2':n_v2,
            'p1':round(n_v1/N_seg*100,1),'p2':round(n_v2/N_seg*100,1)})
    R['transgtipos']=tipos
    return R

# ── Gráficos ──────────────────────────────────────────────────────────────────
def g_sexo(R):
    fig,ax=plt.subplots(figsize=(5,3.5))
    vals=[R['n_hombre'],R['n_mujer']]
    bars=ax.bar(['Hombre','Mujer'],vals,color=[MC_MID,MC_ACCENT],width=0.5,zorder=3)
    for bar,val in zip(bars,vals):
        pct=round(val/R['nv_sex']*100,1) if R['nv_sex']>0 else 0
        ax.text(bar.get_x()+bar.get_width()/2,bar.get_height()+0.3,
                f'{val}\n({pct}%)',ha='center',va='bottom',fontsize=10,fontweight='bold',color='#333')
    ax.set_ylim(0,max(vals)*1.3 if max(vals)>0 else 1)
    ax.set_ylabel('N personas',fontsize=9,color='#595959')
    ax.tick_params(labelsize=9); _ax_style(ax); fig.patch.set_facecolor('white'); fig.tight_layout()
    return fig

def g_edad(R):
    fig,ax=plt.subplots(figsize=(6,4))
    labs=list(R['edad_dist'].keys()); vals=list(R['edad_dist'].values())
    cols=[MC_MID if v==max(vals) else MC_LIGHT for v in vals]
    bars=ax.barh(labs,vals,color=cols,zorder=3)
    for bar,val in zip(bars,vals):
        if val==0: continue
        pct=round(val/R['nv_edad']*100,1) if R['nv_edad']>0 else 0
        ax.text(bar.get_width()+0.1,bar.get_y()+bar.get_height()/2,
                f'{val} ({pct}%)',va='center',fontsize=8,color='#333')
    ax.set_xlim(0,max(vals)*1.5 if max(vals)>0 else 1)
    ax.tick_params(labelsize=9); _ax_style(ax,horiz=True)
    fig.patch.set_facecolor('white'); fig.tight_layout(); return fig

def g_sust_comp(R):
    datos=R['sust_comp']
    if not datos: return None
    labs=[d['label'] for d in datos]; p1=[d['p1'] for d in datos]; p2=[d['p2'] for d in datos]
    x=np.arange(len(labs)); ww=0.35
    fig,ax=plt.subplots(figsize=(max(5.5,len(labs)*1.0),4))
    b1=ax.bar(x-ww/2,p1,ww,color=C_T1,label='Ingreso (TOP 1)',zorder=3)
    b2=ax.bar(x+ww/2,p2,ww,color=C_T2,label='Seguimiento (TOP 2)',zorder=3)
    for bar,v in zip(list(b1)+list(b2),p1+p2):
        if v>0:
            ax.text(bar.get_x()+bar.get_width()/2,bar.get_height()+0.4,
                    f'{v}%',ha='center',va='bottom',fontsize=8,fontweight='bold',color='#333')
    ax.set_xticks(x); ax.set_xticklabels(labs,fontsize=8.5,rotation=15,ha='right')
    ax.set_ylabel('% de personas',fontsize=8,color='#595959')
    ax.set_ylim(0,max(p1+p2)*1.35 if p1+p2 else 1)
    ax.legend(fontsize=8,frameon=False); _ax_style(ax)
    fig.patch.set_facecolor('white'); fig.tight_layout(); return fig

def g_dias_comp(R):
    datos=R['dias_comp']
    if not datos: return None
    labs=[d['label'] for d in datos]; m1=[d['m1'] for d in datos]; m2=[d['m2'] for d in datos]
    x=np.arange(len(labs)); ww=0.35
    fig,ax=plt.subplots(figsize=(max(5,len(labs)*1.0),4))
    b1=ax.bar(x-ww/2,m1,ww,color=C_T1,zorder=3)
    b2=ax.bar(x+ww/2,m2,ww,color=C_T2,zorder=3)
    for bar,v in zip(list(b1)+list(b2),m1+m2):
        if v>0:
            ax.text(bar.get_x()+bar.get_width()/2,bar.get_height()+0.1,
                    f'{v}d',ha='center',va='bottom',fontsize=8,fontweight='bold',color='#333')
    ax.set_xticks(x); ax.set_xticklabels(labs,fontsize=8.5)
    ax.set_ylabel('Promedio días (0–28)',fontsize=8,color='#595959')
    ax.set_ylim(0,max(m1+m2)*1.32 if m1+m2 else 1)
    ax.legend([mpatches.Patch(color=C_T1),mpatches.Patch(color=C_T2)],
              ['Ingreso (TOP 1)','Seguimiento (TOP 2)'],fontsize=8,frameon=False)
    _ax_style(ax); fig.patch.set_facecolor('white'); fig.tight_layout(); return fig

def g_cambio(R):
    datos=R['cambio']
    if not datos: return None
    labs=[d['label'] for d in datos]
    abs_=[d['pct_abs'] for d in datos]; dis=[d['pct_dis'] for d in datos]
    sc_=[d['pct_sc'] for d in datos]; emp=[d['pct_emp'] for d in datos]
    x=np.arange(len(labs))
    fig,ax=plt.subplots(figsize=(max(5.5,len(labs)*1.0),4))
    ax.bar(x,abs_,color=C_ABS2,label='Abstinencia',zorder=3)
    ax.bar(x,dis,bottom=abs_,color=C_DIS2,label='Disminuyó',zorder=3)
    ax.bar(x,sc_,bottom=[a+d for a,d in zip(abs_,dis)],color=C_SC2,label='Sin cambio',zorder=3)
    ax.bar(x,emp,bottom=[a+d+s for a,d,s in zip(abs_,dis,sc_)],color=C_EMP2,label='Empeoró',zorder=3)
    for i,(a,d,s,e) in enumerate(zip(abs_,dis,sc_,emp)):
        y_pos=0
        for val,col in [(a,C_ABS2),(d,C_DIS2),(s,C_SC2),(e,C_EMP2)]:
            if val>9: ax.text(i,y_pos+val/2,f'{val:.0f}%',ha='center',va='center',fontsize=7.5,color='white',fontweight='bold')
            y_pos+=val
    ax.set_xticks(x); ax.set_xticklabels(labs,fontsize=9)
    ax.set_ylabel('% de consumidores al ingreso',fontsize=8,color='#595959'); ax.set_ylim(0,115)
    ax.legend(loc='upper right',fontsize=7.5,frameon=False,ncol=2)
    _ax_style(ax); fig.patch.set_facecolor('white'); fig.tight_layout(); return fig

def g_transgresion(R):
    N=R['N_seg']; cats=['Ingreso (TOP 1)','Seguimiento (TOP 2)']
    pcts=[R['pct_tr1'],R['pct_tr2']]; ns=[R['n_tr1'],R['n_tr2']]
    fig,ax=plt.subplots(figsize=(5,3.5))
    bars=ax.bar(cats,pcts,color=[C_T1,C_T2],width=0.5,zorder=3)
    for bar,pct,n in zip(bars,pcts,ns):
        ax.text(bar.get_x()+bar.get_width()/2,bar.get_height()+0.8,
                f'{pct}%\n(n={n})',ha='center',va='bottom',fontsize=10,fontweight='bold',color='#333')
    ax.set_ylim(0,max(pcts)*1.4 if pcts else 1)
    ax.set_ylabel(f'% sobre N={N}',fontsize=8,color='#595959')
    ax.tick_params(labelsize=9.5); _ax_style(ax)
    fig.patch.set_facecolor('white'); fig.tight_layout(); return fig

def g_tipos_tr(R):
    datos=R['transgtipos']
    if not datos: return None
    labs=[d['label'] for d in datos]; p1=[d['p1'] for d in datos]; p2=[d['p2'] for d in datos]
    x=np.arange(len(labs)); ww=0.35
    fig,ax=plt.subplots(figsize=(max(5,len(labs)*0.9),3.5))
    b1=ax.bar(x-ww/2,p1,ww,color=C_T1,zorder=3)
    b2=ax.bar(x+ww/2,p2,ww,color=C_T2,zorder=3)
    for bar,v in zip(list(b1)+list(b2),p1+p2):
        if v>0:
            ax.text(bar.get_x()+bar.get_width()/2,bar.get_height()+0.3,
                    f'{v}%',ha='center',va='bottom',fontsize=8,fontweight='bold',color='#333')
    ax.set_xticks(x); ax.set_xticklabels(labs,fontsize=8.5)
    ax.set_ylim(0,max(p1+p2)*1.38 if p1+p2 else 1)
    ax.set_ylabel(f'% sobre N={R["N_seg"]}',fontsize=8,color='#595959')
    ax.legend([mpatches.Patch(color=C_T1),mpatches.Patch(color=C_T2)],
              ['Ingreso','Seguimiento'],fontsize=8,frameon=False)
    _ax_style(ax); fig.patch.set_facecolor('white'); fig.tight_layout(); return fig

def g_salud(R):
    datos=R['salud']
    if not datos: return None
    labs=[d['label'] for d in datos]; m1=[d['m1'] for d in datos]; m2=[d['m2'] for d in datos]
    y=np.arange(len(labs)); ww=0.35
    fig,ax=plt.subplots(figsize=(6,3.5))
    b1=ax.barh(y-ww/2,m1,ww,color=C_T1,zorder=3)
    b2=ax.barh(y+ww/2,m2,ww,color=C_T2,zorder=3)
    for bar,v in zip(list(b1)+list(b2),m1+m2):
        ax.text(bar.get_width()+0.2,bar.get_y()+bar.get_height()/2,
                f'{v}',va='center',fontsize=9,fontweight='bold',color='#333')
    ax.set_yticks(y); ax.set_yticklabels(labs,fontsize=9)
    ax.set_xlim(0,26); ax.axvline(x=10,color='#BFBFBF',linestyle='--',linewidth=0.8)
    ax.set_xlabel('Promedio (0–20)',fontsize=8,color='#595959')
    ax.legend([mpatches.Patch(color=C_T1),mpatches.Patch(color=C_T2)],
              ['Ingreso (TOP 1)','Seguimiento (TOP 2)'],fontsize=8,frameon=False)
    _ax_style(ax,horiz=True); fig.patch.set_facecolor('white'); fig.tight_layout(); return fig

def g_vivienda(R):
    cats=['Lugar estable','Condiciones básicas']
    p1=[R['viv1_t1'][1],R['viv2_t1'][1]]; p2=[R['viv1_t2'][1],R['viv2_t2'][1]]
    x=np.arange(len(cats)); ww=0.35
    fig,ax=plt.subplots(figsize=(5,3.5))
    b1=ax.bar(x-ww/2,p1,ww,color=C_T1,zorder=3)
    b2=ax.bar(x+ww/2,p2,ww,color=C_T2,zorder=3)
    for bar,v in zip(list(b1)+list(b2),p1+p2):
        ax.text(bar.get_x()+bar.get_width()/2,bar.get_height()+1.0,
                f'{v}%',ha='center',va='bottom',fontsize=9,fontweight='bold',color='#333')
    ax.set_xticks(x); ax.set_xticklabels(cats,fontsize=10)
    ax.set_ylim(0,118); ax.set_ylabel('% personas con condición',fontsize=8,color='#595959')
    ax.legend([mpatches.Patch(color=C_T1),mpatches.Patch(color=C_T2)],
              ['Ingreso (TOP 1)','Seguimiento (TOP 2)'],fontsize=8,frameon=False)
    _ax_style(ax); fig.patch.set_facecolor('white'); fig.tight_layout(); return fig

# ── Helper salto de página limpio ─────────────────────────────────────────────
def _page_break(doc):
    """Salto de página sin párrafo vacío visible — va dentro del último párrafo de contenido."""
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)

# ── Construcción del Word ─────────────────────────────────────────────────────
def build_word(R):
    doc=Document()
    for sec in doc.sections:
        sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
        sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)

    N=R['N_seg']; pct_seg=round(N/R['N_total']*100,1) if R['N_total']>0 else 0

    # ── Portada ───────────────────────────────────────────────────────────────
    tbl=doc.add_table(rows=1,cols=1); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    c=tbl.rows[0].cells[0]; set_cell_bg(c,C_NAVY)
    for txt,sz,bold in [
        ('INFORME DE SEGUIMIENTO',18,True),
        ('Monitoreo de Resultados de Tratamiento — Instrumento TOP',11,False),
        ('Comparativo Ingreso (TOP1) vs Seguimiento (TOP2)',11,False),
        (NOMBRE_SERVICIO.upper(),14,True),(PERIODO,10,False),
    ]:
        p=c.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=p.add_run(txt); run.font.name='Calibri'; run.font.size=Pt(sz)
        run.font.bold=bold; run.font.color.rgb=WHITE
    doc.add_paragraph()

    # ── KPIs ──────────────────────────────────────────────────────────────────
    _st=R.get('seg_tiempo',{})
    _tv=f'{_st["mediana"]} m' if _st.get('mediana') is not None else '—'
    add_kpi_row(doc,[
        (R['N_total'],'Total ingresaron'),
        (N,f'Con seguimiento ({pct_seg}%)'),
        (R['edad_media'],'Edad promedio'),
        (_tv,'Mediana seguimiento'),
    ])

    # ── Presentación ──────────────────────────────────────────────────────────
    add_section_header(doc,'','Presentación')
    add_body(doc,
        f'Este informe compara los resultados de {N} personas en {NOMBRE_SERVICIO} ({PERIODO}), '
        f'entre el ingreso (TOP 1) y el seguimiento (TOP 2). '
        f'Del total de {R["N_total"]} que ingresaron, {N} ({pct_seg}%) cuentan con ambas evaluaciones. '
        f'La sustancia principal al ingreso fue {R["sust_top1"]} ({R["sust_top1_pct"]}%).')
    if _st.get('mediana'):
        add_body(doc,
            f'El tiempo entre TOP1 y TOP2 fue en mediana de {_st["mediana"]} meses '
            f'(rango: {_st["min"]}–{_st["max"]} meses; N={_st["n"]}).')

    # ── Sección 1: Antecedentes ───────────────────────────────────────────────
    _page_break(doc)
    add_section_header(doc,'1','Antecedentes Generales')

    add_subsection(doc,'1.1. Distribución según Sexo')
    buf,w=fig_to_img(g_sexo(R),12)
    add_picture_kwnext(doc,buf,w)
    add_body(doc,f'De las {N} personas con seguimiento, {R["n_hombre"]} ({R["pct_hombre"]}%) son hombres y {R["n_mujer"]} ({R["pct_mujer"]}%) son mujeres.')
    add_note(doc,f'N válido: {R["nv_sex"]} personas.')

    add_subsection(doc,'1.2. Distribución según Edad')
    buf,w=fig_to_img(g_edad(R),12)
    add_picture_kwnext(doc,buf,w)
    rm=max(R['edad_dist'],key=R['edad_dist'].get) if R['edad_dist'] else '—'
    nm=R['edad_dist'].get(rm,0); pm=round(nm/R['nv_edad']*100,1) if R['nv_edad']>0 else 0
    add_body(doc,f'Promedio de edad: {R["edad_media"]} años (DE={R["edad_sd"]}; rango {R["edad_min"]}–{R["edad_max"]}). Rango más frecuente: {rm} ({nm} personas, {pm}%).')
    add_note(doc,f'N válido: {R["nv_edad"]} personas.')

    # ── Sección 2: Consumo ────────────────────────────────────────────────────
    _page_break(doc)
    add_section_header(doc,'2','Consumo de Sustancias')

    add_subsection(doc,'2.1. Sustancia Principal — Ingreso vs Seguimiento')
    fig_sc=g_sust_comp(R)
    if fig_sc:
        buf,w=fig_to_img(fig_sc,13)
        add_picture_kwnext(doc,buf,w)
        sc=R['sust_comp']
        top2_d=max(sc,key=lambda x:x['n2']) if sc else {'label':'—','p2':0}
        add_body(doc,f'Al ingreso la sustancia principal fue {R["sust_top1"]} ({R["sust_top1_pct"]}%). Al seguimiento: {top2_d["label"]} ({top2_d["p2"]}%).')
        add_note(doc,f'N válido: TOP1={R["nv_sust1"]}, TOP2={R["nv_sust2"]}.')

    add_subsection(doc,'2.2. Promedio de Días de Consumo — Ingreso vs Seguimiento')
    fig_dc=g_dias_comp(R)
    if fig_dc:
        buf,w=fig_to_img(fig_dc,13)
        add_picture_kwnext(doc,buf,w)
        add_body(doc,'Promedio de días de consumo en las últimas 4 semanas al ingreso y al seguimiento (sobre N total, incluyendo 0). Se espera reducción entre TOP1 y TOP2.')
        add_note(doc,f'N: {N} pacientes con ambas evaluaciones.')

    add_subsection(doc,'2.3. Cambio en el Consumo por Sustancia')
    fig_cb=g_cambio(R)
    if fig_cb:
        buf,w=fig_to_img(fig_cb,13)
        add_picture_kwnext(doc,buf,w)
        c=R['cambio']
        pct_abst=round(sum(d['pct_abs'] for d in c)/len(c),1) if c else 0
        add_body(doc,f'En promedio, el {pct_abst}% de los consumidores de cada sustancia logró abstinencia al seguimiento.')
        add_note(doc,'% calculado sobre consumidores al ingreso (días > 0 en TOP1).')

    # ── Sección 3: Transgresión ───────────────────────────────────────────────
    _page_break(doc)
    add_section_header(doc,'3','Transgresión a la Norma Social')

    add_subsection(doc,'3.1. Transgresión — Ingreso vs Seguimiento')
    buf,w=fig_to_img(g_transgresion(R),12)
    add_picture_kwnext(doc,buf,w)
    reduc=round(R['pct_tr1']-R['pct_tr2'],1)
    add_body(doc,
        f'Al ingreso {R["n_tr1"]} personas ({R["pct_tr1"]}%) cometieron alguna transgresión. '
        f'Al seguimiento: {R["n_tr2"]} ({R["pct_tr2"]}%). '
        f'Reducción de {reduc} puntos porcentuales.')
    add_note(doc,f'N total: {N} pacientes.')

    fig_tt=g_tipos_tr(R)
    if fig_tt:
        add_subsection(doc,'3.2. Tipos de Transgresión — Ingreso vs Seguimiento')
        buf,w=fig_to_img(fig_tt,13)
        add_picture_kwnext(doc,buf,w)
        add_body(doc,'Los % no suman 100% (una persona puede cometer más de un tipo).')
        add_note(doc,f'N base: {N} pacientes.')

    # ── Sección 4: Salud y Vivienda ───────────────────────────────────────────
    # Sin salto forzado: keepNext en add_section_header ya evita que quede solo
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_section_header(doc,'4','Salud, Calidad de Vida y Vivienda')

    if R['salud']:
        add_subsection(doc,'4.1. Autopercepción del Estado de Salud (0–20)')
        fig_sal=g_salud(R)
        if fig_sal:
            buf,w=fig_to_img(fig_sal,13)
            add_picture_kwnext(doc,buf,w)
            mejor=max(R['salud'],key=lambda s:s['m2'])
            add_body(doc,f'Mayor mejora en {mejor["label"]} ({flecha(mejor["m1"],mejor["m2"],True)}). Escala 0–20; línea punteada = punto medio (10).')
            add_note(doc,f'N: TOP1={R["salud"][0]["nv1"]}, TOP2={R["salud"][0]["nv2"]}.')

    add_subsection(doc,'4.2. Condiciones de Vivienda — Ingreso vs Seguimiento')
    fig_viv=g_vivienda(R)
    if fig_viv:
        buf,w=fig_to_img(fig_viv,12)
        add_picture_kwnext(doc,buf,w)
        n1_1,p1_1,_=R['viv1_t1']; n1_2,p1_2,_=R['viv1_t2']
        n2_1,p2_1,_=R['viv2_t1']; n2_2,p2_2,_=R['viv2_t2']
        add_body(doc,
            f'Lugar estable: {p1_1}% al ingreso → {p1_2}% al seguimiento ({flecha(p1_1,p1_2,True)}). '
            f'Condiciones básicas: {p2_1}% al ingreso → {p2_2}% al seguimiento ({flecha(p2_1,p2_2,True)}).')
        add_note(doc,f'N: {N} pacientes.')

    # ── Pie ───────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    p=doc.add_paragraph()
    run=p.add_run(f'Informe generado automáticamente · TOP · {NOMBRE_SERVICIO} · {PERIODO}')
    run.font.size=Pt(8); run.font.italic=True; run.font.color.rgb=GRAY
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUTPUT_FILE)
    print(f'  ✓ Word generado: {OUTPUT_FILE}')

if __name__=='__main__':
    print('='*60); print('  SCRIPT_TOP_Universal_Word_Seguimiento  —  Iniciando...'); print('='*60)
    R=cargar_datos()
    print(f'  N_total={R["N_total"]} | N_seg={R["N_seg"]} | {R["sust_top1"]} {R["sust_top1_pct"]}%')
    build_word(R)
    print(f'\n{"="*60}'); print(f'  ✅  LISTO  →  {OUTPUT_FILE}'); print(f'{"="*60}')
